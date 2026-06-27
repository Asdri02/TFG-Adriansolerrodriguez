"""
web/app.py — FastAPI app que expone el grader, el OCR, el dashboard
de validación y el "Aula del Profesor" (asistente IA, calibración few-shot,
sinónimos y antipatrones).

Lanzar desde la raíz del proyecto:
    PYTHONPATH=src .venv_mac/bin/uvicorn web.app:app --reload --port 8000
"""

from __future__ import annotations

import json
import re
import statistics as stats
import sys
import tempfile
from collections import defaultdict
from dataclasses import asdict
from pathlib import Path
from typing import Any, Dict, List, Optional

_SRC = Path(__file__).parent.parent
if str(_SRC) not in sys.path:
    sys.path.insert(0, str(_SRC))

# Carga variables de entorno desde .env en la raíz del proyecto (si existe).
# Necesario para que ANTHROPIC_API_KEY y demás secretos lleguen al cliente de
# Anthropic. Se hace ANTES de importar reference_db / anthropic.
try:
    from dotenv import load_dotenv
    load_dotenv(_SRC.parent / ".env")
except ImportError:
    pass

from fastapi import FastAPI, File, Form, HTTPException, UploadFile
from fastapi.responses import FileResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel
from pytesseract import TesseractNotFoundError

from ai import answer_checker
from ai.models import ReferenceAnswer
from ai.semantic_grader import SemanticGrader
from ocr.extractor import OCRExtractor
from ocr.parser import ExamTextParser
from validate import TEST_CASES
from web import calibration_db
from web import gradebook_db
from web import templates_db
import reference_db

_STATIC_DIR = Path(__file__).parent / "static"
_DATA_DIR = _SRC.parent / "data"
_TEACHER_CONFIG_PATH = _DATA_DIR / "teacher_config.json"

app = FastAPI(title="TFG Exam Grader", docs_url="/api/docs")
app.mount("/static", StaticFiles(directory=str(_STATIC_DIR)), name="static")

# Modelo de Claude usado por todos los endpoints que llaman a la API. Centralizado
# para poder cambiarlo en un solo sitio (texto + Vision).
_CLAUDE_MODEL = "claude-haiku-4-5"


def _resolve_tesseract_cmd() -> Optional[str]:
    """
    Localiza el binario de Tesseract aunque no esté en el PATH del proceso
    Python (típico al lanzar uvicorn desde una terminal vieja tras instalar
    Tesseract con Homebrew). Sigue este orden:
      1. PATH del proceso.
      2. Rutas estándar conocidas en macOS (brew arm64/x86) y Linux.
    Devuelve None si no encuentra nada; pytesseract entonces fallará con su
    error habitual.
    """
    import shutil
    cmd = shutil.which("tesseract")
    if cmd:
        return cmd
    candidates = [
        "/opt/homebrew/bin/tesseract",   # macOS Apple Silicon
        "/usr/local/bin/tesseract",      # macOS Intel
        "/usr/bin/tesseract",            # Linux
    ]
    for path in candidates:
        if Path(path).exists():
            return path
    return None


_extractor = OCRExtractor(tesseract_cmd=_resolve_tesseract_cmd())


# ── Schemas ──────────────────────────────────────────────────────────────────

class KeyConcept(BaseModel):
    concept: str
    weight: float
    # Paráfrasis que el profesor acepta como equivalentes al término
    # (p.ej. "último en entrar primero en salir" para "LIFO"). Opcional.
    synonyms: List[str] = []


class BonusTerm(BaseModel):
    term: str
    weight: float


class ReferencePayload(BaseModel):
    question: str
    subject: str = "General"
    education_level: str = "Bachillerato"
    ideal_answer: str
    key_concepts: List[KeyConcept]
    bonus_terms: List[BonusTerm] = []


class GradeRequest(BaseModel):
    student_answer: str
    reference: ReferencePayload


class GradeCaseRequest(BaseModel):
    case_id: int
    student_answer: Optional[str] = None


class BatchAnswer(BaseModel):
    id: Optional[str] = None  # alumno o identificador, opcional
    text: str


class GradeBatchRequest(BaseModel):
    reference: ReferencePayload
    answers: List[BatchAnswer]


class GenerateReferenceRequest(BaseModel):
    question: str
    subject: str = "General"
    education_level: str = "Bachillerato"
    force: bool = False  # invalida caché y vuelve a llamar a Claude


class CalibrationExample(BaseModel):
    answer: str
    score: float  # 0-10, la que el profesor pondría


class CalibrateRequest(BaseModel):
    reference: ReferencePayload
    student_answer: str
    examples: List[CalibrationExample]


class SynonymGroup(BaseModel):
    canonical: str
    variants: List[str]


class AntiPattern(BaseModel):
    concept: str
    forbidden: List[str]
    penalty: float = 0.5  # multiplicador aplicado al score si se detecta


class TeacherConfig(BaseModel):
    synonyms: List[SynonymGroup] = []
    antipatterns: List[AntiPattern] = []


# ── Helpers: serialization ───────────────────────────────────────────────────

def _reference_from_payload(payload: ReferencePayload) -> ReferenceAnswer:
    return ReferenceAnswer(
        question=payload.question,
        subject=payload.subject,
        education_level=payload.education_level,
        expected_answer_type="respuesta_corta",
        ideal_answer=payload.ideal_answer,
        key_concepts=[c.model_dump() for c in payload.key_concepts],
        bonus_terms=[b.model_dump() for b in payload.bonus_terms],
    )


def _serialize_reference(ref: ReferenceAnswer) -> Dict[str, Any]:
    data = asdict(ref)
    data["rubric"] = data.get("rubric") or []
    data["common_mistakes"] = data.get("common_mistakes") or []
    return data


def _serialize_case(case: Dict[str, Any]) -> Dict[str, Any]:
    ref: ReferenceAnswer = case["reference"]
    return {
        "id": case["id"],
        "topic": case["topic"],
        "desc": case["desc"],
        "subject": ref.subject,
        "education_level": ref.education_level,
        "question": ref.question,
        "ideal_answer": ref.ideal_answer,
        "key_concepts": list(ref.key_concepts),
        "student_answer": case["student_answer"],
        "nota_min": case["nota_min"],
        "nota_max": case["nota_max"],
        "expected_to_fail": case.get("expected_to_fail", False),
    }


# ── Helpers: teacher_config ──────────────────────────────────────────────────

def _load_teacher_config() -> Dict[str, Any]:
    if _TEACHER_CONFIG_PATH.exists():
        try:
            with open(_TEACHER_CONFIG_PATH, "r", encoding="utf-8") as fh:
                return json.load(fh)
        except (json.JSONDecodeError, OSError):
            return {"synonyms": [], "antipatterns": []}
    return {"synonyms": [], "antipatterns": []}


def _save_teacher_config(cfg: Dict[str, Any]) -> None:
    _TEACHER_CONFIG_PATH.parent.mkdir(parents=True, exist_ok=True)
    with open(_TEACHER_CONFIG_PATH, "w", encoding="utf-8") as fh:
        json.dump(cfg, fh, ensure_ascii=False, indent=2)


def _make_grader_with_teacher() -> SemanticGrader:
    """SemanticGrader con el synonym_map extendido con los sinónimos del profesor."""
    grader = SemanticGrader()
    cfg = _load_teacher_config()
    for group in cfg.get("synonyms", []):
        canonical = group.get("canonical", "").strip()
        variants = [v.strip() for v in group.get("variants", []) if v.strip()]
        if not canonical or not variants:
            continue
        key = canonical.lower()
        existing = grader.synonym_map.get(key, [])
        merged = sorted(set(existing) | set(variants) | {canonical})
        grader.synonym_map[key] = merged
    return grader


def _apply_antipatterns(
    result: Dict[str, Any],
    student_answer: str,
    reference_concepts: List[Dict[str, Any]],
    teacher_config: Dict[str, Any],
) -> Dict[str, Any]:
    """
    Si la respuesta del alumno contiene una frase prohibida asociada a un
    concepto presente en la rúbrica, multiplica el score por el penalty.
    Solo afecta si ese concepto aparece en la referencia activa.

    La comparación se hace sobre texto normalizado (sin tildes ni puntuación),
    igual que el resto del grader: así un antipatrón escrito "es el nucleo"
    salta también contra "es el núcleo".
    """
    answer_norm = _normalize_for_compare(student_answer)
    ref_concepts_lower = {_normalize_for_compare(c["concept"]) for c in reference_concepts}

    hits: List[Dict[str, Any]] = []
    score = result["score_over_10"]

    for ap in teacher_config.get("antipatterns", []):
        concept = ap.get("concept", "").strip()
        if not concept or _normalize_for_compare(concept) not in ref_concepts_lower:
            continue
        for forbidden in ap.get("forbidden", []):
            f = _normalize_for_compare(forbidden)
            if not f:
                continue
            if f in answer_norm:
                penalty = float(ap.get("penalty", 0.5))
                score = round(score * penalty, 2)
                hits.append({
                    "concept": concept,
                    "phrase": forbidden,
                    "penalty": penalty,
                })
                break  # un hit por antipatrón

    result["score_over_10"] = score
    result["antipatterns_hit"] = hits
    return result


def _grade(student_answer: str, reference: ReferenceAnswer) -> Dict[str, Any]:
    grader = _make_grader_with_teacher()
    result = grader.grade(student_answer, reference)
    cfg = _load_teacher_config()
    result = _apply_antipatterns(result, student_answer, list(reference.key_concepts), cfg)
    result["reference"] = _serialize_reference(reference)
    return result


# ── Endpoints: básicos (igual que antes) ─────────────────────────────────────

@app.get("/")
def index() -> FileResponse:
    return FileResponse(_STATIC_DIR / "index.html")


@app.post("/api/grade")
def grade(req: GradeRequest) -> Dict[str, Any]:
    reference = _reference_from_payload(req.reference)
    return _grade(req.student_answer, reference)


@app.get("/api/cases")
def list_cases() -> List[Dict[str, Any]]:
    return [_serialize_case(c) for c in TEST_CASES]


@app.post("/api/grade_case")
def grade_case(req: GradeCaseRequest) -> Dict[str, Any]:
    case = next((c for c in TEST_CASES if c["id"] == req.case_id), None)
    if case is None:
        raise HTTPException(status_code=404, detail=f"Caso {req.case_id} no encontrado")
    answer = req.student_answer if req.student_answer is not None else case["student_answer"]
    return _grade(answer, case["reference"])


@app.post("/api/ocr")
async def ocr(image: UploadFile = File(...), lang: str = Form("spa")) -> Dict[str, Any]:
    if image.content_type is None or not image.content_type.startswith("image/"):
        raise HTTPException(status_code=400, detail="El archivo debe ser una imagen.")

    suffix = Path(image.filename or "upload").suffix or ".png"
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        tmp.write(await image.read())
        tmp_path = Path(tmp.name)

    try:
        raw_text = _extractor.extract_text_from_image(str(tmp_path), lang=lang)
    except TesseractNotFoundError:
        raise HTTPException(
            status_code=503,
            detail=(
                "Tesseract no está instalado en el sistema. "
                "Instálalo con `brew install tesseract tesseract-lang` "
                "para usar esta función."
            ),
        )
    finally:
        tmp_path.unlink(missing_ok=True)

    parsed = ExamTextParser.parse_question_and_answer(raw_text)
    return {
        "raw_text": raw_text,
        "question": parsed["question"],
        "student_answer": parsed["student_answer"],
    }


@app.get("/api/validate")
def validate_endpoint() -> Dict[str, Any]:
    grader = _make_grader_with_teacher()
    cfg = _load_teacher_config()

    per_topic: Dict[str, Dict[str, int]] = defaultdict(
        lambda: {"pass": 0, "expected_fail": 0, "unexpected_fail": 0, "total": 0}
    )
    per_subject: Dict[str, Dict[str, int]] = defaultdict(
        lambda: {"pass": 0, "expected_fail": 0, "unexpected_fail": 0, "total": 0}
    )

    cases_out: List[Dict[str, Any]] = []
    passed = expected_fails = unexpected_fails = 0

    for case in TEST_CASES:
        ref: ReferenceAnswer = case["reference"]
        result = grader.grade(case["student_answer"], ref)
        result = _apply_antipatterns(result, case["student_answer"], list(ref.key_concepts), cfg)
        score = result["score_over_10"]
        ok = case["nota_min"] <= score <= case["nota_max"]
        expected_to_fail = case.get("expected_to_fail", False)

        if ok:
            status = "pass"; passed += 1
        elif expected_to_fail:
            status = "expected_fail"; expected_fails += 1
        else:
            status = "unexpected_fail"; unexpected_fails += 1

        topic = case["topic"]; subject = ref.subject
        per_topic[topic][status] += 1; per_topic[topic]["total"] += 1
        per_subject[subject][status] += 1; per_subject[subject]["total"] += 1

        cases_out.append({
            "id": case["id"], "topic": topic, "subject": subject, "desc": case["desc"],
            "score": score,
            "nota_min": case["nota_min"], "nota_max": case["nota_max"],
            "status": status,
        })

    total = len(TEST_CASES)
    conformant = passed + expected_fails

    return {
        "total": total, "passed": passed,
        "expected_fails": expected_fails, "unexpected_fails": unexpected_fails,
        "conformant_pct": round(conformant / total * 100, 1) if total else 0.0,
        "per_topic": dict(per_topic), "per_subject": dict(per_subject),
        "cases": cases_out,
    }


@app.get("/api/correlation")
def correlation_endpoint() -> Dict[str, Any]:
    """
    Validación por CONCORDANCIA: corrige el conjunto `gold_dataset.GOLD`
    (respuestas con nota humana de referencia) con el grader determinista y
    devuelve Spearman/Pearson/MAE/RMSE, el desglose por pregunta y los pares
    (nota_humana, nota_sistema) para dibujar el diagrama de dispersión.

    Las notas humanas son de referencia (asignadas por el autor según rúbrica),
    no de un tribunal real: la métrica mide concordancia con un criterio
    explícito, no exactitud absoluta.
    """
    from ai import metrics
    from ai.gold_dataset import GOLD

    grader = _make_grader_with_teacher()
    cfg = _load_teacher_config()

    all_sys: List[float] = []
    all_hum: List[float] = []
    points: List[Dict[str, Any]] = []
    per_question: List[Dict[str, Any]] = []

    for reference, samples in GOLD:
        sys_scores, hum_scores = [], []
        for answer, human in samples:
            result = grader.grade(answer, reference)
            result = _apply_antipatterns(result, answer, list(reference.key_concepts), cfg)
            s = result["score_over_10"]
            sys_scores.append(s)
            hum_scores.append(float(human))
            points.append({
                "subject": reference.subject,
                "question": reference.question,
                "human": float(human),
                "system": s,
                "answer": answer,
            })
        all_sys += sys_scores
        all_hum += hum_scores
        per_question.append({
            "subject": reference.subject,
            "question": reference.question,
            "n": len(samples),
            "spearman": round(metrics.spearman(sys_scores, hum_scores), 4),
        })

    report = metrics.correlation_report(all_sys, all_hum)
    return {
        **report,
        "per_question": per_question,
        "points": points,
        "note": (
            "Notas humanas de referencia (asignadas por el autor según rúbrica), "
            "no de un tribunal real. Mide concordancia con un criterio explícito."
        ),
    }


# ── Endpoints: corrección por lotes ──────────────────────────────────────────

@app.post("/api/grade_batch")
def grade_batch(req: GradeBatchRequest) -> Dict[str, Any]:
    if not req.answers:
        raise HTTPException(status_code=400, detail="No has enviado respuestas.")
    reference = _reference_from_payload(req.reference)
    grader = _make_grader_with_teacher()
    cfg = _load_teacher_config()

    results: List[Dict[str, Any]] = []
    scores: List[float] = []

    for i, ans in enumerate(req.answers):
        text = ans.text.strip()
        if not text:
            continue
        r = grader.grade(text, reference)
        r = _apply_antipatterns(r, text, list(reference.key_concepts), cfg)
        results.append({
            "id": ans.id or f"#{i + 1}",
            "answer": text,
            "score": r["score_over_10"],
            "concept_ratio": r["concept_ratio"],
            "similarity_ratio": r["similarity_ratio"],
            "length_penalty": r["length_penalty"],
            "detected": r["detected_concepts"],
            "partial": r["partial_concepts"],
            "missing": r["missing_concepts"],
            "antipatterns_hit": r.get("antipatterns_hit", []),
        })
        scores.append(r["score_over_10"])

    if not scores:
        raise HTTPException(status_code=400, detail="Todas las respuestas estaban vacías.")

    # Distribución por bandas para histograma
    bands = [0, 1, 2, 3, 4, 5, 6, 7, 8, 9, 10]
    histogram = [0] * (len(bands) - 1)
    for s in scores:
        idx = min(int(s), len(histogram) - 1)
        histogram[idx] += 1

    return {
        "results": results,
        "stats": {
            "count": len(scores),
            "mean": round(stats.mean(scores), 2),
            "median": round(stats.median(scores), 2),
            "stdev": round(stats.stdev(scores), 2) if len(scores) > 1 else 0.0,
            "min": round(min(scores), 2),
            "max": round(max(scores), 2),
            "pass_count": sum(1 for s in scores if s >= 5),
            "fail_count": sum(1 for s in scores if s < 5),
        },
        "histogram": {
            "labels": [f"{bands[i]}-{bands[i+1]}" for i in range(len(histogram))],
            "values": histogram,
        },
        "reference": _serialize_reference(reference),
    }


# ── Endpoints: Aula del Profesor ─────────────────────────────────────────────

@app.post("/api/generate_reference")
def generate_reference(req: GenerateReferenceRequest) -> Dict[str, Any]:
    """Pide a Claude que proponga ideal_answer + key_concepts para una pregunta."""
    if not req.question.strip():
        raise HTTPException(status_code=400, detail="Falta la pregunta.")
    try:
        if req.force:
            reference_db.invalidate(req.question)
        ref = reference_db.get_reference(
            req.question, subject=req.subject, education_level=req.education_level,
        )
    except Exception as exc:
        raise HTTPException(
            status_code=502,
            detail=f"Error llamando a Claude: {type(exc).__name__}: {exc}",
        )
    return {
        "question": ref.question,
        "subject": ref.subject,
        "education_level": ref.education_level,
        "ideal_answer": ref.ideal_answer,
        "key_concepts": list(ref.key_concepts),
        "common_mistakes": list(ref.common_mistakes),
        "confidence": ref.confidence,
    }


@app.post("/api/calibrate_grade")
def calibrate_grade(req: CalibrateRequest) -> Dict[str, Any]:
    """
    Corrige student_answer de dos formas y las devuelve juntas:
      - score_grader: la nota del grader determinista actual (con teacher_config)
      - score_llm: la nota que predice Claude usando los ejemplos del profesor
        como contexto few-shot
    """
    if not req.examples:
        raise HTTPException(
            status_code=400,
            detail="Necesito al menos un ejemplo (respuesta + nota) para calibrar.",
        )

    reference = _reference_from_payload(req.reference)
    determ = _grade(req.student_answer, reference)

    # Few-shot prompt
    examples_block = ""
    for i, ex in enumerate(req.examples, start=1):
        examples_block += (
            f"\nEJEMPLO {i}:\n"
            f"Respuesta del alumno: «{ex.answer.strip()}»\n"
            f"Nota del profesor: {ex.score}\n"
        )

    system_prompt = (
        "Eres un asistente de corrección académica. Vas a aprender el criterio "
        "de un profesor concreto a partir de ejemplos y aplicarlo a una nueva "
        "respuesta. Devuelves ÚNICAMENTE JSON válido, sin bloques de código."
    )
    user_prompt = (
        f"PREGUNTA: {reference.question}\n"
        f"RESPUESTA IDEAL: {reference.ideal_answer}\n"
        f"\nCONCEPTOS CLAVE PONDERADOS:\n"
        + "\n".join(
            f"  - {c['concept']} (peso {c['weight']})" for c in reference.key_concepts
        )
        + f"\n\nEste profesor puntúa así:\n{examples_block}"
        + f"\nAhora puntúa esta NUEVA respuesta siguiendo el mismo criterio:\n"
        f"«{req.student_answer.strip()}»\n\n"
        f"Devuelve JSON con esta estructura exacta:\n"
        f'{{"score": <float 0 a 10>, "reasoning": "<una frase breve justificando>"}}'
    )

    try:
        import anthropic
        client = anthropic.Anthropic()
        response = client.messages.create(
            model=_CLAUDE_MODEL,
            temperature=0,
            max_tokens=400,
            system=system_prompt,
            messages=[{"role": "user", "content": user_prompt}],
        )
        raw = response.content[0].text.strip()
        # quitar fences si acaso
        if raw.startswith("```"):
            raw = re.sub(r"^```[a-z]*\n?", "", raw)
            raw = re.sub(r"\n?```$", "", raw).strip()
        llm = json.loads(raw)
        llm_score = float(llm["score"])
        llm_reason = llm.get("reasoning", "")
    except Exception as exc:
        raise HTTPException(
            status_code=502,
            detail=f"Error llamando a Claude para calibrar: {type(exc).__name__}: {exc}",
        )

    return {
        "score_grader": determ["score_over_10"],
        "score_llm": round(llm_score, 2),
        "delta": round(llm_score - determ["score_over_10"], 2),
        "reasoning": llm_reason,
        "examples_used": len(req.examples),
        "grader_detail": {
            "detected": determ["detected_concepts"],
            "missing": determ["missing_concepts"],
            "partial": determ["partial_concepts"],
            "antipatterns_hit": determ.get("antipatterns_hit", []),
        },
    }


@app.post("/api/calibrate_deterministic")
def calibrate_deterministic(req: CalibrateRequest) -> Dict[str, Any]:
    """
    Calibración DETERMINISTA (sin LLM): aprende un mapeo monótono
    nota_grader → nota_profesor a partir de los ejemplos (respuesta + nota) y lo
    aplica a la nueva respuesta. Corrige el desfase de escala/severidad del grader
    respecto al criterio del profesor, sin alterar el orden de las respuestas.

    Necesita al menos 2 ejemplos. Funciona sin ANTHROPIC_API_KEY.
    """
    from ai.calibration import ScoreCalibrator

    if len(req.examples) < 2:
        raise HTTPException(
            status_code=400,
            detail="La calibración determinista necesita al menos 2 ejemplos "
                   "(respuesta + nota) para aprender el mapeo.",
        )

    reference = _reference_from_payload(req.reference)

    # Nota cruda del grader para cada ejemplo y para la nueva respuesta.
    raw_examples = [_grade(ex.answer, reference)["score_over_10"] for ex in req.examples]
    teacher_scores = [float(ex.score) for ex in req.examples]

    method = "isotonic" if len(req.examples) >= 4 else "linear"
    calibrator = ScoreCalibrator(method=method).fit(raw_examples, teacher_scores)

    determ_new = _grade(req.student_answer, reference)
    raw_new = determ_new["score_over_10"]
    calibrated = calibrator.transform(raw_new)

    fit_rows = [
        {"answer": ex.answer, "teacher": float(ex.score),
         "raw": raw, "calibrated": calibrator.transform(raw)}
        for ex, raw in zip(req.examples, raw_examples)
    ]
    # Error medio sobre los propios ejemplos, antes y después (informativo).
    mae_before = round(sum(abs(r - t) for r, t in zip(raw_examples, teacher_scores)) / len(teacher_scores), 2)
    mae_after = round(sum(abs(calibrator.transform(r) - t) for r, t in zip(raw_examples, teacher_scores)) / len(teacher_scores), 2)

    return {
        "score_grader": raw_new,
        "score_calibrated": calibrated,
        "delta": round(calibrated - raw_new, 2),
        "method": method,
        "mapping": calibrator.describe(),
        "examples_used": len(req.examples),
        "mae_examples_before": mae_before,
        "mae_examples_after": mae_after,
        "fit": fit_rows,
        "grader_detail": {
            "detected": determ_new["detected_concepts"],
            "missing": determ_new["missing_concepts"],
            "negated": determ_new.get("negated_concepts", []),
        },
        "note": (
            "Mapeo monótono: no cambia el ORDEN de las respuestas, solo ajusta la "
            "escala al criterio del profesor. Reproduce su nivel de exigencia: si "
            "sus notas de ejemplo son generosas, la calibración también lo será."
        ),
    }


@app.get("/api/teacher_config")
def get_teacher_config() -> Dict[str, Any]:
    return _load_teacher_config()


@app.post("/api/teacher_config")
def set_teacher_config(cfg: TeacherConfig) -> Dict[str, Any]:
    payload = cfg.model_dump()
    _save_teacher_config(payload)
    return payload


# ── Endpoints: Explicar nota con IA ──────────────────────────────────────────

class ExplainRequest(BaseModel):
    reference: ReferencePayload
    student_answer: str
    grade_result: Dict[str, Any]


@app.post("/api/explain_grade")
def explain_grade(req: ExplainRequest) -> Dict[str, Any]:
    """Pide a Claude una explicación en lenguaje natural de la nota."""
    reference = _reference_from_payload(req.reference)
    result = req.grade_result
    if "score_over_10" not in result:
        raise HTTPException(
            status_code=400,
            detail="grade_result debe incluir 'score_over_10' (el resultado de /api/grade).",
        )

    concepts_str = "\n".join(
        f"  - {c['concept']} (peso {c['weight']})" for c in reference.key_concepts
    )
    detected = ", ".join(result.get("detected_concepts", [])) or "—"
    partial = ", ".join(result.get("partial_concepts", [])) or "—"
    missing = ", ".join(result.get("missing_concepts", [])) or "—"
    antipatterns = result.get("antipatterns_hit", [])
    ap_str = "; ".join(
        f"«{h['phrase']}» (penalización ×{h['penalty']})" for h in antipatterns
    ) or "ninguno"

    system_prompt = (
        "Eres un profesor amable explicando a un alumno por qué ha sacado una nota. "
        "Sé breve (3-5 frases), específico, y didáctico. No inventes datos: usa "
        "solo lo que te paso. Habla en segunda persona ('tu respuesta...')."
    )
    user_prompt = (
        f"PREGUNTA: {reference.question}\n"
        f"RESPUESTA IDEAL: {reference.ideal_answer}\n"
        f"\nCONCEPTOS DE LA RÚBRICA:\n{concepts_str}\n"
        f"\nRESPUESTA DEL ALUMNO: «{req.student_answer.strip()}»\n"
        f"\nNOTA OBTENIDA: {result['score_over_10']}/10\n"
        f"  - Conceptos detectados: {detected}\n"
        f"  - Detectados parcialmente: {partial}\n"
        f"  - Conceptos faltantes: {missing}\n"
        f"  - Antipatrones detectados: {ap_str}\n"
        f"\nExplica en 3-5 frases por qué la nota es esa, qué hiciste bien y qué "
        f"te faltó. No repitas literalmente la respuesta ideal — referencia los "
        f"conceptos."
    )

    try:
        import anthropic
        client = anthropic.Anthropic()
        response = client.messages.create(
            model=_CLAUDE_MODEL,
            temperature=0,
            max_tokens=400,
            system=system_prompt,
            messages=[{"role": "user", "content": user_prompt}],
        )
        explanation = response.content[0].text.strip()
    except Exception as exc:
        raise HTTPException(
            status_code=502,
            detail=f"Error llamando a Claude: {type(exc).__name__}: {exc}",
        )

    return {"explanation": explanation}


# ── Endpoints: Export texto a DOCX ───────────────────────────────────────────

class ExportDocxRequest(BaseModel):
    question: Optional[str] = None
    answer: Optional[str] = None
    raw_text: Optional[str] = None
    title: str = "Examen transcrito"


@app.post("/api/export_docx")
def export_docx(req: ExportDocxRequest):
    """Genera un .docx con el texto extraído por OCR."""
    from docx import Document  # carga perezosa
    from fastapi.responses import StreamingResponse
    import io

    doc = Document()
    doc.add_heading(req.title, level=1)

    if req.question:
        doc.add_heading("Pregunta", level=2)
        doc.add_paragraph(req.question)

    if req.answer:
        doc.add_heading("Respuesta", level=2)
        doc.add_paragraph(req.answer)

    if req.raw_text and not (req.question or req.answer):
        doc.add_heading("Texto extraído", level=2)
        doc.add_paragraph(req.raw_text)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": 'attachment; filename="examen.docx"'},
    )


# ── Endpoints: Calibración SQLite (BD de ejemplos) ───────────────────────────

class CalibrationExamplePayload(BaseModel):
    question: str
    subject: str = ""
    answer: str
    score: float


@app.get("/api/calibration/examples")
def calibration_list(question: Optional[str] = None) -> List[Dict[str, Any]]:
    return calibration_db.list_examples(question)


@app.get("/api/calibration/questions")
def calibration_questions() -> List[Dict[str, Any]]:
    return calibration_db.list_questions()


@app.post("/api/calibration/examples")
def calibration_add(payload: CalibrationExamplePayload) -> Dict[str, Any]:
    try:
        return calibration_db.add_example(
            payload.question, payload.subject, payload.answer, payload.score
        )
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))


@app.delete("/api/calibration/examples/{example_id}")
def calibration_delete(example_id: int) -> Dict[str, Any]:
    ok = calibration_db.delete_example(example_id)
    if not ok:
        raise HTTPException(status_code=404, detail="Ejemplo no encontrado")
    return {"deleted": example_id}


# ── Endpoints: Mates/Física experimental (paso a paso vía Claude) ────────────

class GradeStepsRequest(BaseModel):
    question: str
    student_answer: str
    subject: str = "Matemáticas"  # o "Física"
    max_points: float = 10.0


@app.post("/api/grade_steps")
def grade_steps(req: GradeStepsRequest) -> Dict[str, Any]:
    """
    ⚠ EXPERIMENTAL. Fuera del alcance core del TFG (rompe interpretabilidad).
    Delega a Claude la corrección paso a paso de un ejercicio cuantitativo:
    valora el procedimiento, propaga errores de apartados previos.
    """
    if not req.question.strip() or not req.student_answer.strip():
        raise HTTPException(status_code=400, detail="Faltan question o student_answer.")

    system_prompt = (
        "Eres un corrector experto de ejercicios cuantitativos de "
        f"{req.subject} a nivel Bachillerato/Universidad. Evalúas paso a paso, "
        "valoras el procedimiento más allá del resultado final, y si en un "
        "apartado posterior se arrastra un error de uno anterior, no penalizas "
        "dos veces — valoras la coherencia del razonamiento (carry-through). "
        "Devuelves ÚNICAMENTE JSON válido."
    )
    user_prompt = (
        f"EJERCICIO:\n{req.question.strip()}\n\n"
        f"RESPUESTA DEL ALUMNO:\n{req.student_answer.strip()}\n\n"
        f"Devuelve JSON con esta estructura exacta:\n"
        f'{{\n'
        f'  "score": <float 0 a {req.max_points}>,\n'
        f'  "steps": [\n'
        f'    {{"name": "<apartado o paso>", "ok": <true|false>, '
        f'"points_obtained": <float>, "points_max": <float>, '
        f'"comment": "<observación breve>"}}\n'
        f'  ],\n'
        f'  "carry_through_note": "<si has aplicado carry-through, explica brevemente; si no, cadena vacía>",\n'
        f'  "summary": "<resumen 1-2 frases>"\n'
        f'}}\n\n'
        f"La suma de points_max debe ser {req.max_points}."
    )

    try:
        import anthropic
        client = anthropic.Anthropic()
        response = client.messages.create(
            model=_CLAUDE_MODEL,
            temperature=0,
            max_tokens=1200,
            system=system_prompt,
            messages=[{"role": "user", "content": user_prompt}],
        )
        raw = response.content[0].text.strip()
        if raw.startswith("```"):
            raw = re.sub(r"^```[a-z]*\n?", "", raw)
            raw = re.sub(r"\n?```$", "", raw).strip()
        data = json.loads(raw)
    except Exception as exc:
        raise HTTPException(
            status_code=502,
            detail=f"Error llamando a Claude: {type(exc).__name__}: {exc}",
        )

    return {
        "score": float(data.get("score", 0.0)),
        "steps": data.get("steps", []),
        "carry_through_note": data.get("carry_through_note", ""),
        "summary": data.get("summary", ""),
        "experimental_warning": (
            "Esta función delega 100% en un LLM y, a diferencia del grader "
            "principal, no es interpretable. Trátala como segunda opinión."
        ),
    }


# ── Endpoints: Examen estructurado (tablas, test, mixto vía Vision) ──────────
#
# La extracción usa Claude Vision (LLM solo para leer el contenido visual);
# la corrección sigue siendo determinista (comparación normalizada + fuzzy +
# sinónimos del profesor). Así separamos "qué pone en la imagen" de "qué nota
# se merece", manteniendo la interpretabilidad del grader.

import base64
from difflib import SequenceMatcher


_EXTRACT_PROMPT = """\
Eres un asistente que extrae el contenido de un examen escolar fotografiado.

Devuelve ÚNICAMENTE un objeto JSON válido (sin markdown, sin texto extra) con
esta estructura:

{
  "type": "table" | "questions",
  "title": "<enunciado principal del ejercicio, si lo hay>",
  "instructions": "<instrucciones adicionales del ejercicio, si las hay>",
  "headers": ["Col1", "Col2", ...],
  "rows": [
    [
      {"text": "<contenido literal de la celda>", "kind": "printed" | "student" | "blank"},
      ...
    ],
    ...
  ]
}

Reglas:
- "printed": texto que estaba impreso originalmente en el examen (parte del
  enunciado, encabezados, datos de contexto pre-cargados).
- "student": texto añadido a mano por el alumno (a boli/lápiz, con caligrafía
  distinta del impreso).
- "blank": celda que el alumno dejó SIN RELLENAR (un hueco, una línea de
  guiones, o simplemente vacía).

Si una celda es ilegible, escribe "[ilegible]" como texto. NO inventes
contenido que no esté en la imagen.

Si no hay tabla (p.ej. son preguntas sueltas), usa type "questions" y omite
headers; cada row será [{"text": "<pregunta>", "kind": "printed"}, {"text":
"<respuesta>", "kind": "student" | "blank"}].
"""


def _call_claude_vision(image_bytes: bytes, media_type: str, prompt: str, *, max_tokens: int = 3000) -> str:
    import anthropic
    client = anthropic.Anthropic()
    b64 = base64.standard_b64encode(image_bytes).decode("ascii")
    response = client.messages.create(
        model=_CLAUDE_MODEL,
        temperature=0,
        max_tokens=max_tokens,
        messages=[{
            "role": "user",
            "content": [
                {"type": "image", "source": {"type": "base64", "media_type": media_type, "data": b64}},
                {"type": "text", "text": prompt},
            ],
        }],
    )
    return response.content[0].text


def _strip_json_fences(raw: str) -> str:
    raw = raw.strip()
    if raw.startswith("```"):
        raw = re.sub(r"^```[a-z]*\n?", "", raw)
        raw = re.sub(r"\n?```$", "", raw).strip()
    return raw


@app.post("/api/extract_structured")
async def extract_structured(image: UploadFile = File(...)) -> Dict[str, Any]:
    """Extrae estructura (tabla o preguntas) de la imagen usando Claude Vision."""
    if image.content_type is None or not image.content_type.startswith("image/"):
        raise HTTPException(status_code=400, detail="El archivo debe ser una imagen.")

    image_bytes = await image.read()
    media_type = image.content_type or "image/png"

    try:
        raw = _call_claude_vision(image_bytes, media_type, _EXTRACT_PROMPT)
        data = json.loads(_strip_json_fences(raw))
    except json.JSONDecodeError as exc:
        raise HTTPException(
            status_code=502,
            detail=f"Claude devolvió JSON inválido: {exc}",
        )
    except Exception as exc:
        raise HTTPException(
            status_code=502,
            detail=f"Error llamando a Claude Vision: {type(exc).__name__}: {exc}",
        )

    # Sanidad mínima
    if "type" not in data or "rows" not in data:
        raise HTTPException(
            status_code=502,
            detail="Estructura inesperada en la respuesta de Claude.",
        )

    return data


class StructuredStructure(BaseModel):
    type: str
    title: Optional[str] = ""
    instructions: Optional[str] = ""
    headers: List[str] = []
    rows: List[List[Dict[str, Any]]]


class GenerateSolutionsRequest(BaseModel):
    structure: StructuredStructure
    subject: str = "General"


@app.post("/api/generate_solutions")
def generate_solutions(req: GenerateSolutionsRequest) -> Dict[str, Any]:
    """
    Pide a Claude la respuesta correcta para cada celda evaluable.

    Importante: enviamos a Claude la estructura CON LAS RESPUESTAS DEL ALUMNO
    ENMASCARADAS (placeholder "[?]"). Si le mandásemos el texto del alumno
    intacto, en muchos casos lo repite asumiendo que es la respuesta correcta
    (bug observado en formulación química con tablas mixtas). Sin el texto del
    alumno, Claude se ve obligado a derivar la respuesta del contexto impreso.
    """
    s = req.structure

    # Esqueleto: las celdas printed mantienen su texto (contexto del enunciado);
    # student/blank se enmascaran a "[?]" para que Claude las trate como huecos.
    skeleton_rows = []
    for row in s.rows:
        new_row = []
        for cell in row:
            if cell.get("kind") == "printed":
                new_row.append({"kind": "printed", "text": cell.get("text", "")})
            else:
                new_row.append({"kind": "to_fill", "text": "[?]"})
        skeleton_rows.append(new_row)

    payload = {
        "type": s.type, "title": s.title, "instructions": s.instructions,
        "headers": s.headers, "rows": skeleton_rows,
    }
    prompt = (
        f"Tienes el ESQUELETO de un ejercicio de {req.subject}. Las celdas con "
        f'kind "printed" son el contexto impreso del examen (enunciado, '
        f"encabezados, datos dados). Las celdas con \"[?]\" son huecos que el "
        f"alumno debía rellenar. Tu tarea: para CADA hueco, dime cuál es la "
        f"respuesta correcta esperada según el contexto.\n\n"
        f"EJERCICIO (esqueleto):\n{json.dumps(payload, ensure_ascii=False, indent=2)}\n\n"
        f"Devuelve ÚNICAMENTE JSON con esta estructura (sin markdown):\n"
        f'{{"solutions": [\n'
        f'  {{"row": <índice 0-based>, "col": <índice 0-based>, "correct": "<respuesta correcta>"}},\n'
        f"  ...\n"
        f"]}}\n\n"
        f"Solo incluye celdas con [?]. Si no puedes determinar la respuesta para "
        f'una celda concreta (p.ej. porque la nomenclatura "Tradicional" no se '
        f'aplica a esa fórmula), omítela. Sé conciso: solo el contenido que '
        f"debería ir en la celda, no explicación."
    )

    try:
        import anthropic
        client = anthropic.Anthropic()
        response = client.messages.create(
            model=_CLAUDE_MODEL,
            temperature=0,
            max_tokens=2000,
            system="Eres un profesor experto. Devuelves SOLO JSON válido.",
            messages=[{"role": "user", "content": prompt}],
        )
        data = json.loads(_strip_json_fences(response.content[0].text))
    except Exception as exc:
        raise HTTPException(
            status_code=502,
            detail=f"Error llamando a Claude: {type(exc).__name__}: {exc}",
        )

    return data


class GradeStructuredRequest(BaseModel):
    structure: StructuredStructure
    solutions: List[Dict[str, Any]]  # [{row, col, correct}]
    fuzzy_threshold: float = 0.80
    points_per_cell: float = 1.0


def _normalize_for_compare(text: str) -> str:
    """Normaliza para comparación: lowercase, sin acentos, sin puntuación trivial."""
    import unicodedata
    t = (text or "").lower().strip()
    t = unicodedata.normalize("NFD", t)
    t = "".join(ch for ch in t if unicodedata.category(ch) != "Mn")
    # NFD + descarte de combining marks ya ha convertido ñ→n, así que no hay
    # tilde de ñ que preservar aquí. Mantenemos paréntesis para nomenclatura
    # del tipo "hierro (III)". \w (modo unicode) cubre letras y dígitos.
    t = re.sub(r"[^\w\s()]", " ", t)
    t = re.sub(r"\s+", " ", t).strip()
    return t


def _apply_synonyms(text: str, synonyms: List[Dict[str, Any]]) -> str:
    """Reemplaza variantes por canónico para que el matching las equipare."""
    out = text
    for g in synonyms:
        canonical = (g.get("canonical") or "").strip()
        if not canonical:
            continue
        for variant in g.get("variants", []):
            v = variant.strip()
            if not v:
                continue
            out = re.sub(rf"\b{re.escape(_normalize_for_compare(v))}\b",
                         _normalize_for_compare(canonical), out)
    return out


def _grade_cells(
    cells_in_structure: List[List[Dict[str, Any]]],
    solutions: List[Dict[str, Any]],
    fuzzy_threshold: float = 0.80,
    points_per_cell: float = 1.0,
) -> Dict[str, Any]:
    """
    Núcleo del grading celda a celda. Reusable para grade_structured y para
    la corrección contra plantilla.

    Para cada solución {row, col, correct}, busca la celda en la estructura,
    extrae el texto del alumno (cualquier kind), compara y devuelve veredicto.
    """
    cfg = _load_teacher_config()
    synonyms = cfg.get("synonyms", [])

    cells_out: List[Dict[str, Any]] = []
    earned = 0.0
    max_points = 0.0

    for sol in solutions:
        try:
            r = int(sol["row"]); c = int(sol["col"])
            correct = str(sol.get("correct", "")).strip()
        except (KeyError, ValueError, TypeError):
            continue
        if r < 0 or r >= len(cells_in_structure): continue
        if c < 0 or c >= len(cells_in_structure[r]): continue
        if not correct: continue

        cell = cells_in_structure[r][c]
        kind = cell.get("kind", "")
        student_text = str(cell.get("text", "")).strip()

        max_points += points_per_cell

        if kind == "blank" or not student_text or student_text in {"---", "----------", "-----------"}:
            verdict = "blank"; points = 0.0
            similarity = 0.0
        else:
            sn = _apply_synonyms(_normalize_for_compare(student_text), synonyms)
            cn = _apply_synonyms(_normalize_for_compare(correct), synonyms)
            if sn == cn:
                verdict = "correct"; points = points_per_cell
                similarity = 1.0
            else:
                similarity = SequenceMatcher(None, sn, cn).ratio()
                if similarity >= fuzzy_threshold:
                    verdict = "partial"; points = points_per_cell * 0.7
                else:
                    verdict = "wrong"; points = 0.0

        earned += points
        cells_out.append({
            "row": r, "col": c,
            "student_text": student_text,
            "correct": correct,
            "verdict": verdict,
            "similarity": round(similarity, 3),
            "points": round(points, 2),
            "points_max": points_per_cell,
        })

    pct = (earned / max_points * 100) if max_points else 0.0
    return {
        "cells": cells_out,
        "earned": round(earned, 2),
        "max_points": round(max_points, 2),
        "score_pct": round(pct, 1),
        "score_over_10": round(pct / 10, 2),
    }


@app.post("/api/grade_structured")
def grade_structured(req: GradeStructuredRequest) -> Dict[str, Any]:
    return _grade_cells(
        req.structure.rows, req.solutions,
        req.fuzzy_threshold, req.points_per_cell,
    )


# ── Endpoints: Plantillas de examen reutilizables ────────────────────────────
#
# Modelo:
#   - exam_templates: estructura + soluciones definidas una vez por el profesor
#   - template_gradings: cada vez que se aplica la plantilla a un alumno,
#     queda guardada la corrección.
# La estructura de cada plantilla:
#   {
#     "type": "table",
#     "title": "...",
#     "headers": [...],
#     "rows": [
#       [
#         {"role": "context", "text": "Fe2S3"},
#         {"role": "evaluable", "correct": "Trisulfuro de dihierro"},
#         {"role": "evaluable", "correct": "Sulfuro de hierro (III)"},
#         {"role": "none"}
#       ], ...
#     ]
#   }

class TemplateCellContext(BaseModel):
    role: str  # "context" | "evaluable" | "none"
    text: Optional[str] = None      # para role=context
    correct: Optional[str] = None   # para role=evaluable


class TemplateStructure(BaseModel):
    type: str = "table"
    title: Optional[str] = ""
    instructions: Optional[str] = ""
    headers: List[str] = []
    rows: List[List[Dict[str, Any]]]


class TemplateCreate(BaseModel):
    name: str
    subject: str = ""
    education_level: str = ""
    structure: TemplateStructure
    points_per_cell: float = 1.0


class TemplateUpdate(BaseModel):
    name: Optional[str] = None
    subject: Optional[str] = None
    education_level: Optional[str] = None
    structure: Optional[TemplateStructure] = None
    points_per_cell: Optional[float] = None


@app.get("/api/templates")
def templates_list() -> List[Dict[str, Any]]:
    return templates_db.list_templates()


@app.post("/api/templates")
def templates_create(payload: TemplateCreate) -> Dict[str, Any]:
    return templates_db.create_template(
        name=payload.name, subject=payload.subject,
        education_level=payload.education_level,
        structure=payload.structure.model_dump(),
        points_per_cell=payload.points_per_cell,
    )


@app.get("/api/templates/{template_id}")
def templates_get(template_id: int) -> Dict[str, Any]:
    t = templates_db.get_template(template_id)
    if not t:
        raise HTTPException(status_code=404, detail="Plantilla no encontrada")
    return t


@app.put("/api/templates/{template_id}")
def templates_update(template_id: int, payload: TemplateUpdate) -> Dict[str, Any]:
    if templates_db.get_template(template_id) is None:
        raise HTTPException(status_code=404, detail="Plantilla no encontrada")
    fields = payload.model_dump(exclude_none=True)
    if "structure" in fields and hasattr(payload.structure, "model_dump"):
        fields["structure"] = payload.structure.model_dump()
    return templates_db.update_template(template_id, **fields)


@app.delete("/api/templates/{template_id}")
def templates_delete(template_id: int) -> Dict[str, Any]:
    ok = templates_db.delete_template(template_id)
    if not ok:
        raise HTTPException(status_code=404, detail="Plantilla no encontrada")
    return {"deleted": template_id}


@app.get("/api/templates/{template_id}/gradings")
def templates_list_gradings(template_id: int) -> List[Dict[str, Any]]:
    return templates_db.list_gradings(template_id)


@app.get("/api/templates/{template_id}/stats")
def templates_stats(template_id: int) -> Dict[str, Any]:
    return templates_db.template_stats(template_id)


@app.delete("/api/templates/{template_id}/gradings/{grading_id}")
def templates_delete_grading(template_id: int, grading_id: int) -> Dict[str, Any]:
    ok = templates_db.delete_grading(grading_id)
    if not ok:
        raise HTTPException(status_code=404, detail="Grading no encontrado")
    return {"deleted": grading_id}


# ── Importar plantilla desde imagen (examen en blanco o ya corregido) ────────

class TemplateFromImagePayload(BaseModel):
    mode: str  # "blank" | "corrected" — qué representa la imagen
    subject: str = "General"
    # nombre opcional, sino se autogenera; el front lo pedirá al guardar
    name: Optional[str] = None


@app.post("/api/templates/from_image")
async def templates_from_image(
    image: UploadFile = File(...),
    mode: str = Form("blank"),
    subject: str = Form("General"),
) -> Dict[str, Any]:
    """
    mode="blank": la imagen es el examen sin respuestas. Extrae estructura y
                  Claude propone las correctas para cada hueco.
    mode="corrected": la imagen es el examen con las respuestas correctas
                  escritas. Extrae todo y usa las del alumno como soluciones.
    Devuelve una estructura LISTA para crear plantilla (no la persiste todavía).
    """
    if mode not in {"blank", "corrected"}:
        raise HTTPException(status_code=400, detail="mode debe ser 'blank' o 'corrected'")
    if image.content_type is None or not image.content_type.startswith("image/"):
        raise HTTPException(status_code=400, detail="El archivo debe ser una imagen.")

    image_bytes = await image.read()
    media_type = image.content_type or "image/png"

    try:
        raw = _call_claude_vision(image_bytes, media_type, _EXTRACT_PROMPT)
        extracted = json.loads(_strip_json_fences(raw))
    except Exception as exc:
        raise HTTPException(
            status_code=502,
            detail=f"Error extrayendo con Vision: {type(exc).__name__}: {exc}",
        )

    # Construir estructura de plantilla
    template_rows: List[List[Dict[str, Any]]] = []
    for row in extracted.get("rows", []):
        new_row = []
        for cell in row:
            kind = cell.get("kind", "")
            text = (cell.get("text") or "").strip()

            if kind == "printed":
                new_row.append({"role": "context", "text": text})
            elif mode == "corrected" and kind == "student" and text and text not in {"---", "----------"}:
                # En modo "corregido", las respuestas del alumno SON las correctas
                new_row.append({"role": "evaluable", "correct": text})
            elif kind in {"blank", "student"}:
                # Hueco. En modo "blank", la solución la propondrá Claude después.
                new_row.append({"role": "evaluable", "correct": ""})
            else:
                new_row.append({"role": "none"})
        template_rows.append(new_row)

    structure = {
        "type": extracted.get("type", "table"),
        "title": extracted.get("title", ""),
        "instructions": extracted.get("instructions", ""),
        "headers": extracted.get("headers", []),
        "rows": template_rows,
    }

    # Si modo="blank", pedimos a Claude también las respuestas correctas
    solutions_warning: Optional[str] = None
    if mode == "blank":
        try:
            # Reusamos generate_solutions pero pasándole el extracted con kinds
            payload = {
                "type": extracted.get("type"), "title": extracted.get("title"),
                "instructions": extracted.get("instructions"),
                "headers": extracted.get("headers"), "rows": extracted.get("rows"),
            }
            req = GenerateSolutionsRequest(
                structure=StructuredStructure(**payload), subject=subject,
            )
            sols = generate_solutions(req).get("solutions", [])
            # Volcar al template_rows
            for s in sols:
                r, c = int(s["row"]), int(s["col"])
                if 0 <= r < len(template_rows) and 0 <= c < len(template_rows[r]):
                    if template_rows[r][c].get("role") == "evaluable":
                        template_rows[r][c]["correct"] = str(s.get("correct", "")).strip()
        except Exception as exc:
            # No abortamos: devolvemos la estructura con las correctas vacías,
            # pero avisamos al front para que el profesor las rellene a mano en
            # vez de creer que la generación funcionó.
            solutions_warning = (
                "No se pudieron autogenerar las respuestas correctas "
                f"({type(exc).__name__}). Revísalas o rellénalas manualmente."
            )

    return {
        "structure": structure,
        "extracted_raw": extracted,
        "solutions_warning": solutions_warning,
    }


# ── Aplicar plantilla a una imagen de alumno ────────────────────────────────

class ApplyTemplateResponse(BaseModel):
    grading_id: int
    student_name: str
    score_over_10: float
    extracted: Dict[str, Any]
    grade_result: Dict[str, Any]


def _structure_to_solutions(template_structure: Dict[str, Any]) -> List[Dict[str, Any]]:
    """De la estructura de plantilla saca la lista de soluciones (row, col, correct)."""
    out = []
    for ri, row in enumerate(template_structure.get("rows", [])):
        for ci, cell in enumerate(row):
            if cell.get("role") == "evaluable":
                correct = (cell.get("correct") or "").strip()
                if correct:
                    out.append({"row": ri, "col": ci, "correct": correct})
    return out


@app.post("/api/templates/{template_id}/grade_image")
async def templates_grade_image(
    template_id: int,
    image: UploadFile = File(...),
    student_name: str = Form(...),
) -> Dict[str, Any]:
    """
    Aplica una plantilla a la foto de un alumno:
      1. Vision extrae lo que escribió el alumno (estructura completa)
      2. Comparamos contra las soluciones guardadas en la plantilla
      3. Guardamos el grading en BD
    """
    template = templates_db.get_template(template_id)
    if not template:
        raise HTTPException(status_code=404, detail="Plantilla no encontrada")
    if not student_name.strip():
        raise HTTPException(status_code=400, detail="Falta student_name")
    if image.content_type is None or not image.content_type.startswith("image/"):
        raise HTTPException(status_code=400, detail="El archivo debe ser una imagen.")

    image_bytes = await image.read()
    media_type = image.content_type or "image/png"

    try:
        raw = _call_claude_vision(image_bytes, media_type, _EXTRACT_PROMPT)
        extracted = json.loads(_strip_json_fences(raw))
    except Exception as exc:
        raise HTTPException(
            status_code=502,
            detail=f"Error extrayendo con Vision: {type(exc).__name__}: {exc}",
        )

    solutions = _structure_to_solutions(template["structure"])
    if not solutions:
        raise HTTPException(
            status_code=400,
            detail="La plantilla no tiene celdas evaluables con respuesta correcta definida.",
        )

    grade = _grade_cells(
        extracted.get("rows", []), solutions,
        fuzzy_threshold=0.80,
        points_per_cell=template.get("points_per_cell", 1.0),
    )

    grading = templates_db.add_grading(
        template_id=template_id,
        student_name=student_name,
        extracted=extracted,
        grade_result=grade,
    )

    return {
        "grading_id": grading["id"],
        "student_name": grading["student_name"],
        "score_over_10": grade["score_over_10"],
        "earned": grade["earned"],
        "max_points": grade["max_points"],
        "extracted": extracted,
        "grade_result": grade,
        "created_at": grading["created_at"],
    }


# ── Cuaderno del profesor: clases, alumnos, exámenes y notas ──────────────────
#
# Flujo pensado para una clase de 30: dar de alta la clase y el listado de
# alumnos, crear un examen con su rúbrica (la misma ReferencePayload del grader)
# y corregir las 30 respuestas de una vez, dejando las notas guardadas y fechadas.

@app.get("/aula")
def aula() -> FileResponse:
    return FileResponse(_STATIC_DIR / "gradebook.html")


class ClassCreate(BaseModel):
    name: str
    subject: str = ""
    academic_year: str = ""


class StudentCreate(BaseModel):
    name: str


class StudentsBulkCreate(BaseModel):
    names: List[str]


class ExamCreate(BaseModel):
    title: str
    subject: str = ""
    exam_date: str = ""
    # Modo de corrección del examen:
    #   conceptual → grader semántico (Bio, Historia, Filosofía, Química...)
    #   numeric    → checker SymPy de resultado (Mates/Física)
    #   writing    → juez LLM con rúbrica (Inglés/Lengua)
    grading_mode: str = "conceptual"
    rubric: Dict[str, Any]


class ClassAnswer(BaseModel):
    student_id: int
    text: str


class GradeClassRequest(BaseModel):
    answers: List[ClassAnswer]


# -- Clases --

@app.get("/api/gradebook/classes")
def gb_list_classes() -> List[Dict[str, Any]]:
    return gradebook_db.list_classes()


@app.post("/api/gradebook/classes")
def gb_create_class(payload: ClassCreate) -> Dict[str, Any]:
    try:
        return gradebook_db.create_class(payload.name, payload.subject, payload.academic_year)
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))


@app.delete("/api/gradebook/classes/{class_id}")
def gb_delete_class(class_id: int) -> Dict[str, Any]:
    if not gradebook_db.delete_class(class_id):
        raise HTTPException(status_code=404, detail="Clase no encontrada")
    return {"deleted": class_id}


# -- Alumnos --

@app.get("/api/gradebook/classes/{class_id}/students")
def gb_list_students(class_id: int) -> List[Dict[str, Any]]:
    if gradebook_db.get_class(class_id) is None:
        raise HTTPException(status_code=404, detail="Clase no encontrada")
    return gradebook_db.list_students(class_id)


@app.post("/api/gradebook/classes/{class_id}/students")
def gb_add_student(class_id: int, payload: StudentCreate) -> Dict[str, Any]:
    try:
        return gradebook_db.add_student(class_id, payload.name)
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))


@app.post("/api/gradebook/classes/{class_id}/students/bulk")
def gb_add_students_bulk(class_id: int, payload: StudentsBulkCreate) -> List[Dict[str, Any]]:
    if gradebook_db.get_class(class_id) is None:
        raise HTTPException(status_code=404, detail="Clase no encontrada")
    return gradebook_db.add_students_bulk(class_id, payload.names)


@app.delete("/api/gradebook/students/{student_id}")
def gb_delete_student(student_id: int) -> Dict[str, Any]:
    if not gradebook_db.delete_student(student_id):
        raise HTTPException(status_code=404, detail="Alumno no encontrado")
    return {"deleted": student_id}


@app.get("/api/gradebook/students/{student_id}/grades")
def gb_student_grades(student_id: int) -> List[Dict[str, Any]]:
    return gradebook_db.grades_for_student(student_id)


# -- Exámenes --

@app.get("/api/gradebook/classes/{class_id}/exams")
def gb_list_exams(class_id: int) -> List[Dict[str, Any]]:
    if gradebook_db.get_class(class_id) is None:
        raise HTTPException(status_code=404, detail="Clase no encontrada")
    return gradebook_db.list_exams(class_id)


@app.post("/api/gradebook/classes/{class_id}/exams")
def gb_create_exam(class_id: int, payload: ExamCreate) -> Dict[str, Any]:
    mode = payload.grading_mode
    rubric = dict(payload.rubric)
    rubric["grading_mode"] = mode

    # Validamos la rúbrica según el modo elegido.
    if mode == "conceptual":
        try:
            ReferencePayload(**{k: v for k, v in rubric.items() if k != "grading_mode"})
        except Exception:
            raise HTTPException(status_code=400, detail="Rúbrica conceptual inválida (pregunta + conceptos clave).")
    elif mode == "numeric":
        if not str(rubric.get("expected", "")).strip():
            raise HTTPException(status_code=400, detail="Falta el resultado esperado ('expected').")
        if rubric.get("kind", "math") not in {"math", "physics"}:
            raise HTTPException(status_code=400, detail="kind debe ser 'math' o 'physics'.")
    elif mode == "writing":
        if not str(rubric.get("question", "")).strip():
            raise HTTPException(status_code=400, detail="Falta el enunciado/tarea ('question').")
    else:
        raise HTTPException(status_code=400, detail="grading_mode debe ser 'conceptual', 'numeric' o 'writing'.")

    try:
        return gradebook_db.create_exam(
            class_id, payload.title, payload.subject, payload.exam_date, rubric=rubric,
        )
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))


@app.get("/api/gradebook/exams/{exam_id}")
def gb_get_exam(exam_id: int) -> Dict[str, Any]:
    exam = gradebook_db.get_exam(exam_id)
    if exam is None:
        raise HTTPException(status_code=404, detail="Examen no encontrado")
    return exam


@app.delete("/api/gradebook/exams/{exam_id}")
def gb_delete_exam(exam_id: int) -> Dict[str, Any]:
    if not gradebook_db.delete_exam(exam_id):
        raise HTTPException(status_code=404, detail="Examen no encontrado")
    return {"deleted": exam_id}


@app.get("/api/gradebook/exams/{exam_id}/grades")
def gb_list_grades(exam_id: int) -> Dict[str, Any]:
    exam = gradebook_db.get_exam(exam_id)
    if exam is None:
        raise HTTPException(status_code=404, detail="Examen no encontrado")
    return {
        "exam": exam,
        "grades": gradebook_db.list_grades(exam_id),
        "stats": gradebook_db.exam_stats(exam_id),
    }


@app.post("/api/gradebook/exams/{exam_id}/grade_class")
def gb_grade_class(exam_id: int, req: GradeClassRequest) -> Dict[str, Any]:
    """
    Corrige de una vez las respuestas de toda la clase contra la rúbrica del
    examen y guarda cada nota. El corazón del "corregir 30 exámenes a la vez".
    """
    exam = gradebook_db.get_exam(exam_id)
    if exam is None:
        raise HTTPException(status_code=404, detail="Examen no encontrado")
    if not req.answers:
        raise HTTPException(status_code=400, detail="No has enviado respuestas.")

    rubric = exam.get("rubric") or {}
    mode = rubric.get("grading_mode", "conceptual")

    # Para modo conceptual preparamos la referencia una sola vez.
    reference = None
    if mode == "conceptual":
        try:
            reference = _reference_from_payload(
                ReferencePayload(**{k: v for k, v in rubric.items() if k != "grading_mode"})
            )
        except Exception:
            raise HTTPException(
                status_code=400,
                detail="El examen no tiene una rúbrica conceptual válida. Edítalo y vuelve a guardar.",
            )

    def _grade_one(text: str):
        """Devuelve (score, detail) según el modo de corrección del examen."""
        if mode == "numeric":
            r = answer_checker.grade_numeric(
                text, str(rubric.get("expected", "")), kind=rubric.get("kind", "math"),
            )
            return r["score_over_10"], {"correct": r["correct"], "extracted": r.get("extracted"),
                                        "expected": r.get("expected"), "detail": r.get("detail")}
        if mode == "writing":
            res = grade_writing(GradeWritingRequest(
                question=rubric.get("question", ""), student_answer=text,
                subject=rubric.get("subject", "lengua"), criteria=rubric.get("criteria"),
            ))
            return res["score_over_10"], {"criteria": res.get("criteria", []),
                                          "feedback": res.get("feedback", "")}
        # conceptual
        result = _grade(text, reference)
        return result["score_over_10"], {
            "detected": result["detected_concepts"], "partial": result["partial_concepts"],
            "missing": result["missing_concepts"], "concept_ratio": result["concept_ratio"],
            "antipatterns_hit": result.get("antipatterns_hit", []),
        }

    valid_ids = {s["id"] for s in gradebook_db.list_students(exam["class_id"])}
    graded = []
    for ans in req.answers:
        if ans.student_id not in valid_ids:
            continue  # alumno de otra clase / inexistente → se ignora
        text = ans.text.strip()
        if not text:
            continue
        score, detail = _grade_one(text)
        gradebook_db.upsert_grade(exam_id, ans.student_id, score, answer=text, detail=detail)
        graded.append({"student_id": ans.student_id, "score": score})

    if not graded:
        raise HTTPException(status_code=400, detail="Ninguna respuesta válida que corregir.")

    return {
        "graded_count": len(graded),
        "grades": gradebook_db.list_grades(exam_id),
        "stats": gradebook_db.exam_stats(exam_id),
    }


# ── Mates / Física: checker determinista de resultado (+ 2ª opinión LLM) ──────
#
# Para preguntas cuya nota depende del RESULTADO (no de los conceptos): se
# extrae la respuesta final y se compara por equivalencia matemática (SymPy) o
# por valor+unidad (física). Opcionalmente se añade la corrección paso a paso
# del LLM como segunda opinión (que sí puede dar crédito al procedimiento).

class GradeNumericRequest(BaseModel):
    student_answer: str
    expected: str
    kind: str = "math"           # "math" | "physics"
    rel_tol: float = 0.05        # tolerancia relativa para física
    question: str = ""           # enunciado, solo para la 2ª opinión LLM
    with_llm_opinion: bool = False


@app.post("/api/grade_numeric")
def grade_numeric_endpoint(req: GradeNumericRequest) -> Dict[str, Any]:
    if req.kind not in {"math", "physics"}:
        raise HTTPException(status_code=400, detail="kind debe ser 'math' o 'physics'.")
    if not req.expected.strip():
        raise HTTPException(status_code=400, detail="Falta el resultado esperado.")

    result = answer_checker.grade_numeric(
        req.student_answer, req.expected, kind=req.kind, rel_tol=req.rel_tol,
    )

    # Segunda opinión opcional del LLM (crédito al procedimiento aunque el
    # número falle). No bloquea: si no hay API key, se anota el error.
    if req.with_llm_opinion:
        try:
            steps = grade_steps(GradeStepsRequest(
                question=req.question or "(enunciado no facilitado)",
                student_answer=req.student_answer,
                subject="Matemáticas" if req.kind == "math" else "Física",
            ))
            result["llm_opinion"] = steps
        except HTTPException as exc:
            result["llm_opinion"] = {"error": exc.detail}

    return result


# ── Inglés / Lengua: juez LLM con rúbrica ─────────────────────────────────────
#
# La redacción/gramática/comentario NO es evaluable de forma determinista con
# las restricciones del proyecto (sin NLP pesado). Se delega en Claude con una
# rúbrica estructurada por criterios. NO interpretable: se marca como tal.

_WRITING_RUBRICS: Dict[str, List[Dict[str, Any]]] = {
    "ingles": [
        {"id": "task", "label": "Logro de la tarea / contenido", "max": 2.5},
        {"id": "grammar", "label": "Gramática y corrección", "max": 2.5},
        {"id": "vocabulary", "label": "Riqueza y precisión léxica", "max": 2.5},
        {"id": "coherence", "label": "Coherencia y cohesión", "max": 2.5},
    ],
    "lengua": [
        {"id": "tema", "label": "Tema y comprensión del texto", "max": 2.5},
        {"id": "estructura", "label": "Estructura y organización", "max": 2.5},
        {"id": "argumentacion", "label": "Argumentación", "max": 2.5},
        {"id": "expresion", "label": "Expresión y corrección", "max": 2.5},
    ],
}


class GradeWritingRequest(BaseModel):
    question: str
    student_answer: str
    subject: str = "lengua"      # "ingles" | "lengua" (o cualquier clave de rúbrica)
    criteria: Optional[List[Dict[str, Any]]] = None  # rúbrica personalizada opcional


@app.post("/api/grade_writing")
def grade_writing(req: GradeWritingRequest) -> Dict[str, Any]:
    if not req.student_answer.strip():
        raise HTTPException(status_code=400, detail="Falta la respuesta del alumno.")

    rubric = req.criteria or _WRITING_RUBRICS.get(req.subject.lower().strip())
    if not rubric:
        raise HTTPException(
            status_code=400,
            detail=f"No hay rúbrica para '{req.subject}'. Envía 'criteria' o usa 'ingles'/'lengua'.",
        )
    total_max = sum(float(c["max"]) for c in rubric)
    criteria_block = "\n".join(f"  - {c['id']} ({c['label']}): hasta {c['max']} puntos" for c in rubric)

    lang_note = ("Evalúa una redacción en INGLÉS de un examen de Bachillerato. "
                 "Valora también la corrección gramatical y la fluidez del inglés."
                 if req.subject.lower().strip() == "ingles" else
                 "Evalúas una respuesta de Lengua Castellana de Bachillerato (comentario/redacción).")

    system_prompt = (
        f"Eres un corrector de exámenes de Bachillerato. {lang_note} "
        "Puntúas cada criterio de forma independiente y justa según la rúbrica. "
        "Devuelves ÚNICAMENTE JSON válido, sin markdown."
    )
    user_prompt = (
        f"ENUNCIADO/TAREA: {req.question}\n\n"
        f"RESPUESTA DEL ALUMNO:\n«{req.student_answer.strip()}»\n\n"
        f"RÚBRICA (puntúa cada criterio de 0 a su máximo):\n{criteria_block}\n\n"
        f"Devuelve JSON con esta estructura exacta:\n"
        f'{{\n'
        f'  "criteria": [{{"id": "<id>", "score": <float>, "max": <float>, "comment": "<breve>"}}],\n'
        f'  "feedback": "<retroalimentación global 2-3 frases>"\n'
        f'}}\n'
        f"La suma de 'score' será la nota sobre {total_max}."
    )

    try:
        import anthropic
        client = anthropic.Anthropic()
        response = client.messages.create(
            model=_CLAUDE_MODEL, temperature=0, max_tokens=800,
            system=system_prompt,
            messages=[{"role": "user", "content": user_prompt}],
        )
        data = json.loads(_strip_json_fences(response.content[0].text))
    except Exception as exc:
        raise HTTPException(
            status_code=502,
            detail=f"Error llamando a Claude: {type(exc).__name__}: {exc}",
        )

    crit = data.get("criteria", [])
    total = round(sum(float(c.get("score", 0)) for c in crit), 2)
    # Escala a 0-10 por si la rúbrica no suma 10
    score_over_10 = round(total / total_max * 10, 2) if total_max else 0.0

    return {
        "criteria": crit,
        "total": total,
        "total_max": total_max,
        "score_over_10": score_over_10,
        "feedback": data.get("feedback", ""),
        "method": "llm_rubric",
        "warning": (
            "Esta corrección la realiza un LLM con una rúbrica; a diferencia del "
            "grader determinista, no es plenamente interpretable."
        ),
    }


class VerifyAnswerRequest(BaseModel):
    student_answer: str
    reference: ReferencePayload


@app.post("/api/verify_answer")
def verify_answer(req: VerifyAnswerRequest) -> Dict[str, Any]:
    """
    Corrige con el grader determinista y, además, pide al LLM una verificación
    de FACTUALIDAD (segunda opinión). No reescribe la nota: si el verificador
    detecta conceptos negados o mal atribuidos que el grader sí contó, marca el
    caso como 'revisar' para que decida el profesor. Mantiene la trazabilidad.
    """
    from ai.verifier import verify_factuality

    reference = _reference_from_payload(req.reference)
    determ = _grade(req.student_answer, reference)

    try:
        verification = verify_factuality(
            question=reference.question,
            ideal_answer=reference.ideal_answer,
            key_concepts=[dict(c) for c in reference.key_concepts],
            student_answer=req.student_answer,
        )
    except RuntimeError as exc:
        raise HTTPException(status_code=502, detail=str(exc))

    # Conceptos que el grader acreditó pero el verificador considera negados o
    # mal atribuidos: ahí es donde lo determinista se queda corto.
    determ_credited = set(determ.get("detected_concepts", []))
    conflicts = [
        c for c in verification["flagged"] if c["concept"] in determ_credited
    ]
    needs_review = bool(verification["contradiction"] or conflicts)

    return {
        "score_grader": determ["score_over_10"],
        "grader_detail": {
            "detected": determ.get("detected_concepts", []),
            "negated": determ.get("negated_concepts", []),
            "missing": determ.get("missing_concepts", []),
        },
        "verification": verification,
        "conflicts": conflicts,
        "needs_review": needs_review,
        "advice": (
            "Revisar a mano: la nota automática puede sobrevalorar la respuesta "
            "(el alumno usa los términos pero los niega o los atribuye mal)."
            if needs_review else
            "El grader y la verificación de factualidad coinciden."
        ),
    }
